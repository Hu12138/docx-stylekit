import base64
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from docx_stylekit import fix_image_paragraphs
from docx_stylekit.tools.image_paragraphs import _paragraph_contains_image
import subprocess
import sys


PNG_1PX = base64.b64decode(
    b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEvwH+0zm6AwAAAABJRU5ErkJggg=="
)


def test_fix_image_paragraph_spacing(tmp_path):
    img_path = tmp_path / "dot.png"
    img_path.write_bytes(PNG_1PX)

    doc_path = tmp_path / "image.docx"
    doc = Document()
    p = doc.add_paragraph("   ")
    doc.add_picture(str(img_path))
    doc.save(doc_path)

    result_path = fix_image_paragraphs(doc_path)
    assert result_path.exists()

    doc_after = Document(str(result_path))
    paragraphs = [p for p in doc_after.paragraphs if _paragraph_contains_image(p)]
    assert paragraphs, "Should detect image paragraph"
    para = paragraphs[0]
    p_pr = para._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    assert spacing is not None
    assert spacing.get(qn("w:lineRule")) == "auto"
    assert spacing.get(qn("w:line")) == "240"
    jc = p_pr.find(qn("w:jc"))
    assert jc is not None and jc.get(qn("w:val")) == "center"
    ind = p_pr.find(qn("w:ind"))
    assert ind is not None
    assert all(ind.get(attr, "0") == "0" for attr in (
        qn("w:firstLine"),
        qn("w:firstLineChars"),
        qn("w:hanging"),
        qn("w:hangingChars"),
        qn("w:left"),
        qn("w:right"),
    ))


def test_cli_fix_images(tmp_path):
    img_path = tmp_path / "dot.png"
    img_path.write_bytes(PNG_1PX)
    doc_path = tmp_path / "input.docx"
    doc = Document()
    doc.add_picture(str(img_path))
    doc.save(doc_path)
    out_path = tmp_path / "output.docx"
    cmd = [
        sys.executable,
        "-m",
        "docx_stylekit.cli",
        "fix-images",
        str(doc_path),
        "-o",
        str(out_path),
    ]
    subprocess.run(cmd, check=True)
    assert out_path.exists()
