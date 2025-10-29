import base64
import subprocess
import sys
from pathlib import Path

from docx import Document

from docx_stylekit import sanitize_docx

PNG_1PX = base64.b64decode(
    b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEvwH+0zm6AwAAAABJRU5ErkJggg=="
)


def build_docs(tmp_path: Path):
    raw = tmp_path / "raw.docx"
    template = tmp_path / "template.docx"
    image = tmp_path / "tiny.png"
    image.write_bytes(PNG_1PX)

    raw_doc = Document()
    raw_doc.add_paragraph("一、测试标题")
    raw_doc.add_paragraph("正文内容")
    raw_doc.add_picture(str(image))
    raw_doc.save(raw)

    template_doc = Document()
    template_doc.add_heading("模板标题", level=1)
    template_doc.save(template)

    return raw, template


def test_sanitize_api(tmp_path):
    raw, template = build_docs(tmp_path)
    output = tmp_path / "sanitized.docx"
    sanitize_docx(raw, template_docx=template, output_path=output)

    doc = Document(str(output))
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert paragraphs[0].style.name == "Heading 1"
    image_paragraph = next(p for p in doc.paragraphs if p._element.xpath('.//w:drawing'))
    assert image_paragraph.style.name == "ImageParagraph"


def test_sanitize_cli(tmp_path):
    raw, template = build_docs(tmp_path)
    output = tmp_path / "sanitized_cli.docx"
    cmd = [
        sys.executable,
        "-m",
        "docx_stylekit.cli",
        "sanitize",
        str(raw),
        "-t",
        str(template),
        "-o",
        str(output),
    ]
    subprocess.run(cmd, check=True)
    assert output.exists() and output.stat().st_size > 0
    doc = Document(str(output))
    first = next(p for p in doc.paragraphs if p.text.strip())
    assert first.style.name == "Heading 1"
