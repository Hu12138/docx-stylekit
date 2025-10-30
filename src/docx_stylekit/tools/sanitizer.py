from __future__ import annotations

import copy
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import xml.etree.ElementTree as ET

from .image_paragraphs import fix_image_paragraph_spacing
from ..utils.io import load_yaml
from ..writer.docx_writer import _apply_table_format
from ..writer.style_store import StyleResolver
from importlib import resources


MANDATORY_STYLES = ["ImageParagraph", "PageNumber", "InfoTable"]

MANUAL_HEADING_PREFIX = re.compile(
    r"""
    ^\s*
    (
        第[一二三四五六七八九十百千万零]+章 |
        [一二三四五六七八九十]+、 |
        （[一二三四五六七八九十]+） |
        \([一二三四五六七八九十]+\) |
        \d+(?:[\.．]\d+)*(?:[\.．、])? |
        \d+\)
    )
    \s*
    """,
    re.VERBOSE,
)


def _load_default_profile() -> dict:
    resource_path = resources.files("docx_stylekit.data").joinpath("default_render_template.yaml")
    with resources.as_file(resource_path) as path:
        return load_yaml(path)


def _copy_docx(src: Path) -> Path:
    tmp_dir = Path(tempfile.mkdtemp())
    dst = tmp_dir / src.name
    shutil.copy2(src, dst)
    return dst


def _replace_part(docx_path: Path, part_name: str, content: bytes):
    tmp_zip = docx_path.with_suffix(".tmp")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_zip, "w", zipfile.ZIP_DEFLATED) as zout:
        replaced = False
        for item in zin.infolist():
            if item.filename == part_name:
                zout.writestr(item, content)
                replaced = True
            else:
                zout.writestr(item, zin.read(item.filename))
        if not replaced:
            zout.writestr(part_name, content)
    docx_path.unlink()
    tmp_zip.rename(docx_path)


def _extract_part(docx_path: Path, part_name: str) -> Optional[bytes]:
    with zipfile.ZipFile(docx_path) as z:
        if part_name in z.namelist():
            return z.read(part_name)
    return None


def _ensure_paragraph_style(
    doc: Document,
    name: str,
    *,
    base_on: str = "Normal",
    font_name: str = "Times New Roman",
    size_pt: float = 16,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    line_spacing: float = 1.0,
) -> None:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        try:
            style.base_style = doc.styles[base_on]
        except KeyError:
            pass

    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    pf = style.paragraph_format
    pf.alignment = align
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)


def _ensure_table_style(doc: Document, name: str):
    try:
        doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        base_style = None
        for candidate in ("Table Grid", "Normal Table"):
            try:
                base_style = doc.styles[candidate]
                break
            except KeyError:
                continue
        if base_style:
            style.base_style = base_style


def _ensure_required_styles(doc: Document, profile: dict, *, allow_override: bool = True):
    styles_inline = profile.get("doc", {}).get("stylesInline", {}) if profile else {}
    resolver = StyleResolver(doc, styles_inline, prefer_json_styles=allow_override)

    type_mapping = {
        "paragraph": "paragraph",
        "character": "character",
        "table": "table",
    }

    for name, spec in styles_inline.items():
        stype = spec.get("type")
        expected = type_mapping.get(stype)
        if not expected:
            continue
        try:
            doc.styles[name]
            exists = True
        except KeyError:
            exists = False
        if not allow_override and exists and spec.get("$override"):
            continue
        resolver.ensure_style(name, expected)

    # 表格相关样式不依赖 Normal，避免继承正文的首行缩进
    for table_style in ("TableBase", "TableCell", "TableHeaderCell"):
        try:
            st = doc.styles[table_style]
        except KeyError:
            continue
        try:
            st.base_style = None
        except AttributeError:
            pass

    # 如果默认配置缺失某些关键 style，则回退创建一个基础版本，避免后续访问失败
    for style_name in MANDATORY_STYLES:
        try:
            doc.styles[style_name]
        except KeyError:
            if style_name == "InfoTable":
                _ensure_table_style(doc, style_name)
            else:
                _ensure_paragraph_style(doc, style_name)


def _detect_heading_pattern(text: str) -> Optional[tuple[int, str, str]]:
    if not text:
        return None
    stripped = text.strip()
    # Numeric patterns like 1., 1.1, 1.1.1
    numeric_match = re.match(r"^(\d+(?:\.\d+)*)[\.\s]", stripped)
    if numeric_match:
        segments = numeric_match.group(1).split('.')
        remainder = stripped[numeric_match.end():]
        return len(segments), "numeric", remainder
    # Chinese numerals
    if re.match(r"^第[一二三四五六七八九十百千]+章", stripped):
        remainder = stripped[2:]
        return 1, "chapter_cn", remainder
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        remainder = stripped[2:]
        return 1, "chinese", remainder
    paren_cn = re.match(r"^（[一二三四五六七八九十]+）", stripped) or re.match(r"^\([一二三四五六七八九十]+\)", stripped)
    if paren_cn:
        remainder = stripped[paren_cn.end():]
        return 2, "paren_cn", remainder
    paren_num = re.match(r"^（\d+）", stripped) or re.match(r"^\(\d+\)", stripped)
    if paren_num:
        remainder = stripped[paren_num.end():]
        return 2, "paren_num", remainder
    return None


TITLE_KEYWORDS = [
    "说明书",
    "方案",
    "规划",
    "报告",
    "手册",
    "计划",
    "分析",
    "项目",
]


def _extract_orig_heading_level(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    m = re.match(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
    if m:
        return int(m.group(1))
    m = re.match(r"标题\s*(\d+)", style_name)
    if m:
        return int(m.group(1))
    return None


def _map_style_name(
    raw_name: str,
    text: str,
    non_empty_index: int,
    previous_level: Optional[int],
    *,
    in_table: bool = False,
    pattern_level: Optional[int] = None,
    pattern_kind: Optional[str] = None,
    pattern_remainder: str = "",
    previous_pattern_level: Optional[int] = None,
    previous_pattern_kind: Optional[str] = None,
) -> str:
    raw_name = raw_name or ""
    normalized = raw_name.lower()
    stripped = text.strip()
    if not stripped:
        return "Normal"
    if in_table:
        return "TableBase"
    original_level = _extract_orig_heading_level(raw_name)
    detected_level = pattern_level
    detected_kind = pattern_kind
    detected_remainder = pattern_remainder or ""
    if detected_level is None:
        detected = _detect_heading_pattern(stripped)
        if detected:
            detected_level, detected_kind, detected_remainder = detected

    if detected_level:
        level = detected_level
        if previous_level is not None:
            if (
                previous_pattern_level is not None
                and detected_level == previous_pattern_level
                and detected_kind == previous_pattern_kind
            ):
                level = previous_level
            elif level < previous_level:
                if (
                    detected_kind == "numeric"
                    and previous_pattern_kind in {"chinese", "chapter_cn", "paren_cn", "paren_num"}
                ):
                    level = min(previous_level + 1, 9)
                elif (
                    previous_pattern_kind is not None
                    and detected_kind == previous_pattern_kind
                ):
                    level = min(previous_level + 1, 9)
                else:
                    level = detected_level
        if original_level is None:
            remainder_stripped = detected_remainder.strip()
            if len(remainder_stripped) >= 20 or any(ch in remainder_stripped for ch in "：:。！？；，、%"):
                return "Normal"
        if original_level:
            if previous_level is not None:
                target = previous_level + 1
                diff_pattern = abs(detected_level - target)
                diff_original = abs(original_level - target)
                level = detected_level if diff_pattern <= diff_original else original_level
            else:
                level = detected_level if detected_level <= original_level else original_level
        return f"Heading {max(1, min(level, 9))}"

    if original_level:
        if previous_level is not None and previous_level >= 1:
            if original_level > previous_level:
                if len(stripped) <= 20:
                    if original_level == previous_level + 1:
                        return f"Heading {original_level}"
                    if original_level > previous_level + 1:
                        return f"Heading {previous_level + 1}"
                if original_level > previous_level + 1:
                    return f"Heading {previous_level + 1}"
        if previous_level is not None and previous_level >= 1:
            target = previous_level + 1
            level = min(original_level, target)
            return f"Heading {level}"
        return f"Heading {original_level}"

    if non_empty_index == 0 and not in_table:
        if len(stripped) <= 80 and any(keyword in stripped for keyword in TITLE_KEYWORDS):
            return "Title"
        if len(stripped) <= 40:
            return "Title"
    if (
        not in_table
        and stripped
        and len(stripped) <= 25
        and any(stripped.endswith(k) for k in TITLE_KEYWORDS)
    ):
        return "Title"
    return "Normal"


def _clear_run_formatting(paragraph):
    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            run._element.remove(rpr)


def _clear_paragraph_formatting(paragraph):
    ppr = paragraph._element.find(qn("w:pPr"))
    if ppr is not None:
        for child in list(ppr):
            if child.tag in {
                qn("w:ind"),
                qn("w:spacing"),
                qn("w:jc"),
                qn("w:keepNext"),
                qn("w:keepLines"),
                qn("w:pageBreakBefore"),
                qn("w:rPr"),
            }:
                ppr.remove(child)


def _iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _load_defined_num_ids(docx_path: Path) -> set[str]:
    with zipfile.ZipFile(docx_path) as z:
        try:
            numbering_xml = z.read("word/numbering.xml")
        except KeyError:
            return set()
    root = ET.fromstring(numbering_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    num_ids = set()
    for num in root.findall("w:num", ns):
        num_id = num.get(f"{{{ns['w']}}}numId")
        if num_id:
            num_ids.add(num_id)
    return num_ids


def _style_numbering_info(style) -> Optional[tuple[str, Optional[str]]]:
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        return None
    num_id_el = numpr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    ilvl_el = numpr.find(qn("w:ilvl"))
    num_id = num_id_el.get(qn("w:val"))
    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else None
    if not num_id:
        return None
    return num_id, ilvl


def _apply_style_numbering(paragraph, num_id: str, assigned_level: Optional[int], default_ilvl: Optional[str] = None):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    else:
        for child in list(numpr):
            numpr.remove(child)

    if assigned_level is not None and assigned_level >= 1:
        ilvl_val = str(max(0, assigned_level - 1))
    elif default_ilvl is not None:
        ilvl_val = default_ilvl
    else:
        ilvl_val = "0"

    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl_val)
    numpr.append(ilvl_el)

    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    numpr.append(num_id_el)


def _remove_paragraph_numbering(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


def _strip_manual_heading_prefix(paragraph):
    text = paragraph.text
    if not text:
        return
    updated = text
    # 迭代剥离前缀，最多处理数层编号模式，防止 1.1.（一） 等组合残留
    for _ in range(5):
        match = MANUAL_HEADING_PREFIX.match(updated)
        if not match:
            break
        remainder = updated[match.end():].lstrip()
        if not remainder:
            break
        updated = remainder
    if updated != text:
        paragraph.text = updated


def sanitize_docx(
    raw_docx: Path,
    template_docx: Optional[Path] = None,
    *,
    output_path: Optional[Path] = None,
) -> Path:
    raw_docx = Path(raw_docx)
    template_docx = Path(template_docx) if template_docx else None

    working_copy = _copy_docx(raw_docx)
    default_profile = _load_default_profile()
    table_defaults = default_profile.get("doc", {}).get("renderDefaults", {}).get("table", {})
    table_format = table_defaults.get("format", {})

    if template_docx and template_docx.exists():
        styles_xml = _extract_part(template_docx, "word/styles.xml")
        if styles_xml:
            _replace_part(working_copy, "word/styles.xml", styles_xml)
        numbering_xml = _extract_part(template_docx, "word/numbering.xml")
        if numbering_xml:
            _replace_part(working_copy, "word/numbering.xml", numbering_xml)

    doc = Document(str(working_copy))
    original_doc = Document(str(raw_docx))
    _ensure_required_styles(doc, default_profile, allow_override=template_docx is None)
    defined_num_ids = _load_defined_num_ids(working_copy)

    paragraphs = list(_iter_paragraphs(doc))
    original_paragraphs = list(_iter_paragraphs(original_doc))
    original_styles = [p.style.name if p.style else "" for p in original_paragraphs]

    non_empty_counter = 0
    previous_heading_level: Optional[int] = None
    previous_pattern_level: Optional[int] = None
    previous_pattern_kind: Optional[str] = None
    for idx, paragraph in enumerate(paragraphs):
        original_style = original_styles[idx] if idx < len(original_styles) else ""
        in_table = paragraph._element.getparent().tag.endswith('tc')
        text_content = paragraph.text
        stripped_text = text_content.strip()
        pattern_info = _detect_heading_pattern(stripped_text) if stripped_text else None
        pattern_level = pattern_info[0] if pattern_info else None
        pattern_kind = pattern_info[1] if pattern_info else None
        pattern_remainder = pattern_info[2] if pattern_info else ""
        style_name = _map_style_name(
            original_style,
            text_content,
            non_empty_counter,
            previous_heading_level,
            in_table=in_table,
            pattern_level=pattern_level,
            pattern_kind=pattern_kind,
            pattern_remainder=pattern_remainder,
            previous_pattern_level=previous_pattern_level,
            previous_pattern_kind=previous_pattern_kind,
        )
        assigned_level = None
        if style_name.startswith("Heading "):
            try:
                assigned_level = int(style_name.split()[1])
            except ValueError:
                assigned_level = None
        if assigned_level is not None and previous_heading_level is not None and assigned_level > previous_heading_level + 1:
            assigned_level = previous_heading_level + 1
            style_name = f"Heading {assigned_level}"
        style_obj = None
        if paragraph._element.xpath(".//w:drawing"):
            style_obj = doc.styles["ImageParagraph"]
            paragraph.style = style_obj
        else:
            try:
                style_obj = doc.styles[style_name]
                paragraph.style = style_obj
            except KeyError:
                try:
                    style_obj = doc.styles[style_name]
                    paragraph.style = style_obj
                except KeyError:
                    style_obj = doc.styles["Normal"]
                    paragraph.style = style_obj
        _clear_run_formatting(paragraph)
        _clear_paragraph_formatting(paragraph)
        if in_table:
            pf = paragraph.paragraph_format
            pf.left_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        numbering_applied = False
        if assigned_level is not None and style_obj is not None:
            numbering_info = _style_numbering_info(style_obj)
            if numbering_info:
                num_id, default_ilvl = numbering_info
                if num_id in defined_num_ids:
                    _apply_style_numbering(paragraph, num_id, assigned_level, default_ilvl)
                    numbering_applied = True
                else:
                    _remove_paragraph_numbering(paragraph)
            else:
                _remove_paragraph_numbering(paragraph)
        if numbering_applied and pattern_level:
            _strip_manual_heading_prefix(paragraph)
        if paragraph.text.strip():
            if style_name.startswith("Heading "):
                try:
                    previous_heading_level = int(style_name.split()[1])
                except ValueError:
                    previous_heading_level = None
                if assigned_level is not None and pattern_level is not None:
                    previous_pattern_level = pattern_level
                    previous_pattern_kind = pattern_kind
                else:
                    previous_pattern_level = None
                    previous_pattern_kind = None
            elif style_name == "Title":
                previous_heading_level = 0
                previous_pattern_level = None
                previous_pattern_kind = None
            non_empty_counter += 1

    try:
        style = doc.styles["InfoTable"]
    except KeyError:
        style = None
    for table in doc.tables:
        if style:
            try:
                table.style = style
            except KeyError:
                pass
        if table_format:
            _apply_table_format(table, table_format)

    doc.save(str(working_copy))
    fix_image_paragraph_spacing(working_copy, working_copy)

    destination = Path(output_path) if output_path else raw_docx
    shutil.copy2(working_copy, destination)
    return destination


def qn(tag: str) -> str:
    from docx.oxml.ns import qn as _qn

    return _qn(tag)
