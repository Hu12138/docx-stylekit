from __future__ import annotations

from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _paragraph_contains_image(paragraph) -> bool:
    element = paragraph._element
    return bool(
        element.xpath(".//w:drawing")
        or element.xpath('.//*[contains(local-name(), "blip")]')
        or element.xpath('.//*[contains(local-name(), "pic")]')
        or element.xpath('.//*[contains(local-name(), "shape")]')
    )


def _iter_all_paragraphs(doc: Document) -> Iterable:
    seen = set()

    def _from_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if id(paragraph) in seen:
                continue
            seen.add(id(paragraph))
            yield paragraph

    for paragraph in _from_paragraphs(doc.paragraphs):
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in _from_paragraphs(cell.paragraphs):
                    yield paragraph


def _remove_leading_whitespace_runs(paragraph):
    whitespace_chars = {" ", "\t", "\r", "\n", "\u3000"}
    p_element = paragraph._p
    for child in list(p_element):
        if not child.tag.endswith("}r"):
            if child.tag.endswith("}drawing"):
                break
            continue

        if child.xpath(".//w:drawing"):
            break

        text_elems = child.xpath(".//w:t")
        tabs = child.xpath(".//w:tab")
        text = "".join(t.text or "" for t in text_elems)
        if text and any(c not in whitespace_chars for c in text):
            break

        has_tabs = bool(tabs)
        is_whitespace = (text == "" and has_tabs) or all(c in whitespace_chars for c in text)

        if is_whitespace:
            p_element.remove(child)
            continue

        break


def fix_image_paragraph_spacing(input_path: Path, output_path: Optional[Path] = None) -> Path:
    doc = Document(str(input_path))
    updated_count = 0

    for paragraph in _iter_all_paragraphs(doc):
        if not _paragraph_contains_image(paragraph):
            continue

        _remove_leading_whitespace_runs(paragraph)

        p_pr = paragraph._p.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)

        for attr in (
            qn("w:line"),
            qn("w:lineRule"),
            qn("w:before"),
            qn("w:after"),
            qn("w:beforeAutospacing"),
            qn("w:afterAutospacing"),
        ):
            if attr in spacing.attrib:
                del spacing.attrib[attr]

        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:line"), "240")

        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            p_pr.append(jc)
        jc.set(qn("w:val"), "center")

        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)

        for attr in (
            qn("w:firstLine"),
            qn("w:firstLineChars"),
            qn("w:hanging"),
            qn("w:hangingChars"),
            qn("w:start"),
            qn("w:left"),
            qn("w:right"),
            qn("w:end"),
        ):
            if attr in ind.attrib:
                del ind.attrib[attr]

        for attr, value in (
            (qn("w:firstLine"), "0"),
            (qn("w:firstLineChars"), "0"),
            (qn("w:hanging"), "0"),
            (qn("w:hangingChars"), "0"),
            (qn("w:start"), "0"),
            (qn("w:left"), "0"),
            (qn("w:right"), "0"),
            (qn("w:end"), "0"),
        ):
            ind.set(attr, value)

        updated_count += 1

    destination = output_path or input_path
    doc.save(str(destination))
    return destination
