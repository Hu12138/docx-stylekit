import os
import json
import yaml
from typing import Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from .style_store import StyleResolver
from .section_utils import apply_section_layout, add_page_number_field, add_toc_field
from ..utils.dicts import deep_merge

def _clear_cell(cell):
    while cell._tc.getchildren():
        cell._tc.remove(cell._tc.getchildren()[-1])

def _set_cell_shading(cell, color_hex: str):
    if not color_hex:
        return
    color = color_hex.lstrip("#").upper()
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)

def _set_cell_border(cell, border: dict):
    if not border:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ("top", "bottom", "left", "right"):
        cfg = border.get(edge)
        if not cfg:
            continue
        el = tc_borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tc_borders.append(el)
        el.set(qn('w:val'), cfg.get("style", "single"))
        if "color" in cfg:
            el.set(qn('w:color'), cfg["color"].lstrip("#").upper())
        if "size" in cfg:
            el.set(qn('w:sz'), str(int(cfg["size"])))

def _apply_cell_vertical_align(cell, align):
    if not align:
        return
    key = str(align).lower()
    mapping = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "center": WD_ALIGN_VERTICAL.CENTER,
        "middle": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
        "both": WD_ALIGN_VERTICAL.BOTH if hasattr(WD_ALIGN_VERTICAL, "BOTH") else None,
    }
    val = mapping.get(key)
    if val is not None:
        cell.vertical_alignment = val

def _apply_table_format(table, fmt: dict):
    if not fmt:
        return
    header_cfg = fmt.get("header")
    if header_cfg:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, header_cfg.get("fill"))
            _set_cell_border(cell, header_cfg.get("border"))
            _apply_cell_vertical_align(cell, header_cfg.get("verticalAlign"))
            for p in cell.paragraphs:
                for run in p.runs:
                    font = run.font
                    if header_cfg.get("bold") is not None:
                        font.bold = bool(header_cfg.get("bold"))
                    if header_cfg.get("color"):
                        font.color.rgb = RGBColor.from_string(header_cfg["color"].lstrip("#"))
    banding = fmt.get("bandedRows")
    if banding:
        alt_cfg = fmt.get("alternate")
        for idx, row in enumerate(table.rows[1:], start=1):
            if idx % 2 == 1 and alt_cfg:
                for cell in row.cells:
                    _set_cell_shading(cell, alt_cfg.get("fill"))
                    _apply_cell_vertical_align(cell, alt_cfg.get("verticalAlign"))
    table_border = fmt.get("tableBorder")
    if table_border:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.append(tbl_pr)
        borders = tbl_pr.find(qn('w:tblBorders'))
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tbl_pr.append(borders)
        for edge in ("top","bottom","left","right","insideH","insideV"):
            cfg = table_border.get(edge)
            if not cfg:
                continue
            el = borders.find(qn(f'w:{edge}'))
            if el is None:
                el = OxmlElement(f'w:{edge}')
                borders.append(el)
            el.set(qn('w:val'), cfg.get("style","single"))
            if "color" in cfg:
                el.set(qn('w:color'), cfg["color"].lstrip("#").upper())
            if "size" in cfg:
                el.set(qn('w:sz'), str(int(cfg["size"])))
    body_cfg = fmt.get("cell") or fmt.get("body")
    if body_cfg:
        has_header = bool(header_cfg) and len(table.rows) > 0
        start_idx = 1 if has_header else 0
        for row in table.rows[start_idx:]:
            for cell in row.cells:
                _apply_cell_vertical_align(cell, body_cfg.get("verticalAlign"))

def _set_table_widths(table, columns_cfg, section):
    if not columns_cfg or section is None:
        return
    try:
        usable = section.page_width - section.left_margin - section.right_margin
    except AttributeError:
        return
    if usable <= 0:
        return
    widths_pct = []
    unspecified = []
    for idx, col_cfg in enumerate(columns_cfg):
        pct = col_cfg.get("widthPct") if isinstance(col_cfg, dict) else None
        if isinstance(pct, (int, float)) and pct > 0:
            widths_pct.append(float(pct))
        else:
            widths_pct.append(None)
            unspecified.append(idx)
    specified_total = sum(p for p in widths_pct if p is not None)
    remaining = max(0.0, 100.0 - specified_total)
    default_pct = remaining / len(unspecified) if unspecified else 0.0
    for idx in unspecified:
        widths_pct[idx] = default_pct
    total_pct = sum(widths_pct)
    if total_pct <= 0:
        return
    scale = 100.0 / total_pct
    widths_pct = [p * scale for p in widths_pct]
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.append(tbl_pr)
    tblW = tbl_pr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tbl_pr.insert(0, tblW)
    tblW.set(qn('w:w'), str(int(usable)))
    tblW.set(qn('w:type'), 'dxa')
    for col_idx, pct in enumerate(widths_pct):
        col_width = int(round(usable * pct / 100.0))
        if col_idx >= len(table.columns):
            break
        table.columns[col_idx].width = col_width
        for cell in table.columns[col_idx].cells:
            cell.width = col_width
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def _write_cell_blocks(cell, blocks, resolver: StyleResolver):
    _clear_cell(cell)
    for b in blocks or []:
        btype = b.get("type")
        if btype == "paragraph":
            st = resolver.ensure_style(b.get("styleRef", "Normal"), "paragraph")
            p = cell.add_paragraph(style=st.name if st else None)
            for r in b.get("runs", []):
                run = p.add_run(r.get("text", ""))
                cstyle = resolver.ensure_style(r.get("charStyleRef"), "character") if r.get("charStyleRef") else None
                if cstyle:
                    run.style = cstyle
        elif btype == "heading":
            level = int(b.get("level", 1))
            st = resolver.ensure_style(b.get("styleRef", f"Heading {level}"), "paragraph")
            cell.add_paragraph(b.get("text", ""), style=st.name if st else None)
        elif btype == "caption":
            st = resolver.ensure_style(b.get("styleRef", "Caption"), "paragraph")
            p = cell.add_paragraph(style=st.name if st else None)
            p.add_run(b.get("text", ""))

def write_blocks(
    doc: Document,
    blocks: list,
    resolver: StyleResolver,
    fail_on_unknown_style: bool = True,
    table_defaults: Optional[dict] = None,
):
    for b in blocks or []:
        # 页面模板调用：新建节 + 应用布局 + 渲染模板内部 blocks
        if "useTemplate" in b:
            tpl_name = b["useTemplate"]
            page_templates = getattr(doc, "_page_templates_cfg", {})  # 在 render_to_docx 中挂上
            tpl = page_templates.get(tpl_name)
            if not tpl:
                raise ValueError(f"useTemplate 指向的页面模板不存在：{tpl_name}")
            # 新建节
            section = doc.add_section()
            apply_section_layout(section, tpl.get("layout"))
            # 局部变量合并：已在 expand_document 合并到 b["variables"]，此处仅透传
            # 渲染模板 blocks
            write_blocks(
                doc,
                tpl.get("blocks", []),
                resolver,
                fail_on_unknown_style,
                table_defaults=table_defaults,
            )
            continue

        btype = b.get("type")
        if btype == "pageBreak":
            p = doc.add_paragraph()
            p.paragraph_format.page_break_before = True
            continue

        if btype == "toc":
            p = doc.add_paragraph()
            doc_toc = getattr(doc, "_toc_levels", [1,3])
            add_toc_field(p, doc_toc)
            continue

        if btype in ("paragraph", "caption"):
            stname = b.get("styleRef", "Normal")
            st = resolver.ensure_style(stname, "paragraph")
            if fail_on_unknown_style and st is None:
                raise ValueError(f"未知样式（段落）：{stname}")
            p = doc.add_paragraph(style=st.name if st else None)
            if b.get("pageBreakBefore"):
                p.paragraph_format.page_break_before = True
            for r in b.get("runs", []):
                run = p.add_run(r.get("text", ""))
                cstyle = resolver.ensure_style(r.get("charStyleRef"), "character") if r.get("charStyleRef") else None
                if cstyle:
                    run.style = cstyle
            if btype == "caption" and not b.get("runs"):
                p.add_run(b.get("text", ""))
            continue

        if btype == "heading":
            level = int(b.get("level", 1))
            stname = b.get("styleRef", f"Heading {level}")
            st = resolver.ensure_style(stname, "paragraph")
            if fail_on_unknown_style and st is None:
                raise ValueError(f"未知样式（标题）：{stname}")
            doc.add_paragraph(b.get("text", ""), style=st.name if st else None)
            continue

        if btype == "list":
            ordered = bool(b.get("ordered", False))
            stname = b.get("styleRef", "Normal")
            st = resolver.ensure_style(stname, "paragraph")
            if fail_on_unknown_style and st is None:
                raise ValueError(f"未知样式（列表段落）：{stname}")
            for it in b.get("items", []):
                p = doc.add_paragraph(style=st.name if st else None)
                for r in it.get("runs", []):
                    run = p.add_run(r.get("text", ""))
                    cstyle = resolver.ensure_style(r.get("charStyleRef"), "character") if r.get("charStyleRef") else None
                    if cstyle:
                        run.style = cstyle
            continue

        if btype == "table":
            defaults = table_defaults or {}
            columns = b.get("columns") or defaults.get("columns", [])
            if columns:
                ncols = len(columns)
            elif b.get("header"):
                ncols = len(b["header"][0]) if b["header"] else 1
            elif b.get("rows"):
                ncols = len(b["rows"][0]) if b["rows"] else 1
            else:
                ncols = 1
            header = b.get("header", [])
            rows = b.get("rows", [])
            total_rows = len(header) + len(rows)
            if total_rows == 0:
                total_rows = 1
            table = doc.add_table(rows=total_rows, cols=ncols)
            style_ref = b.get("styleRef") or defaults.get("styleRef")
            tstyle = resolver.ensure_style(style_ref, "table") if style_ref else None
            if tstyle:
                table.style = tstyle
            r_idx = 0
            for hrow in header:
                for c_idx, cell in enumerate(hrow):
                    _write_cell_blocks(table.cell(r_idx, c_idx), cell.get("blocks", []), resolver)
                r_idx += 1
            for drow in rows:
                for c_idx, cell in enumerate(drow):
                    _write_cell_blocks(table.cell(r_idx, c_idx), cell.get("blocks", []), resolver)
                r_idx += 1
            current_section = doc.sections[-1] if doc.sections else None
            _set_table_widths(table, columns, current_section)
            table_format = deep_merge(defaults.get("format", {}), b.get("format", {})) if defaults else b.get("format")
            _apply_table_format(table, table_format)
            continue

        # 其它类型（figure 等）可按需扩展
def _clear_document_body(doc: Document):
    """移除模板中的现有正文内容，仅保留节属性。"""
    body = doc._body._element
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


# src/docx_stylekit/writer/docx_writer.py （续）

def render_to_docx(template_json: dict,
                   template_docx_path: Optional[str] = None,
                   styles_yaml: dict = None,
                   output_path: str = "output.docx",
                   prefer_json_styles: bool = False,
                   fail_on_unknown_style: bool = True,
                   clear_existing_content: bool = True):
    """
    template_json: expand_document() 的结果（已展开变量/循环/条件；保留 useTemplate）
    template_docx_path: 样式/编号/页眉页脚基础骨架。可为空（使用内置空白文档）
    styles_yaml: 合并后的 YAML（dict）。如传入路径字符串则会自动读取。
    """
    doc_cfg = template_json.get("doc", {})
    # 读取 JSON 内联样式 / 页面模板 / TOC 级别
    styles_inline = doc_cfg.get("stylesInline", {}) or {}
    page_templates = doc_cfg.get("pageTemplates", {}) or {}
    toc_levels = doc_cfg.get("toc", {}).get("levels", [1,3])

    # 打开模板 DOCX（若未提供，则使用空白文档）
    if template_docx_path:
        doc = Document(template_docx_path)
        if clear_existing_content:
            _clear_document_body(doc)
    else:
        doc = Document()
    # 挂载配置供 writer 使用
    doc._page_templates_cfg = page_templates
    doc._toc_levels = toc_levels

    # 全局 pageSetup（第一节）
    ps = doc_cfg.get("pageSetup", {})
    if ps:
        first = doc.sections[0]
        layout = {
            "marginsCm": ps.get("marginsCm"),
            "orientation": ps.get("orientation"),
            "titleFirstPageDifferent": ps.get("titleFirstPageDifferent"),
            "evenOddDifferent": ps.get("evenOddDifferent"),
            "startAt": ps.get("pageNumbering", {}).get("startAt"),
        }
        apply_section_layout(first, layout)

    # YAML 样式库：此实现依赖模板 DOCX 自带样式；YAML 用于校验/提示（如需从 YAML 动态创建，可在此扩展）
    if isinstance(styles_yaml, str) and os.path.exists(styles_yaml):
        with open(styles_yaml, "r", encoding="utf-8") as f:
            _ = yaml.safe_load(f)

    # 构建解析器（支持 JSON 动态新增样式 & 受控覆盖）
    resolver = StyleResolver(doc, styles_inline, prefer_json_styles=prefer_json_styles)
    # 预加载所有内联样式，确保字体/颜色覆盖立即生效
    for name, style_def in styles_inline.items():
        stype = style_def.get("type")
        if stype in ("paragraph", "character", "table"):
            try:
                resolver.ensure_style(name, stype)
            except ValueError:
                # table 样式在文档缺失时跳过，由后续调用按需创建
                continue

    # TOC（顶层 doc.toc.required 也可以在 blocks 里单独放 type:"toc" 控制位置）
    # 若需要固定在文档开头：可以在此插入；这里尊重 blocks 中的显式位置。

    # 内容写入
    table_defaults = doc_cfg.get("renderDefaults", {}).get("table", {})
    write_blocks(
        doc,
        doc_cfg.get("blocks", []),
        resolver,
        fail_on_unknown_style=fail_on_unknown_style,
        table_defaults=table_defaults,
    )

    # 页眉页脚页码（若 JSON 指定 pageNumber 项，模板未内置时可插入）
    hf = doc_cfg.get("headersFooters", {})
    if "footer" in hf:
        for comp in hf["footer"]:
            if comp.get("type") == "pageNumber":
                for section in doc.sections:
                    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
                    style_ref = comp.get("styleRef")
                    if style_ref:
                        style_obj = resolver.ensure_style(style_ref, "paragraph")
                        if style_obj:
                            fp.style = style_obj
                    if not fp.text.strip():
                        add_page_number_field(fp, align=comp.get("align", "center"))

    doc.save(output_path)
